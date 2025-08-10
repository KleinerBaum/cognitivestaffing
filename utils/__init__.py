"""Utility helpers for the Vacalyser app."""

import os
import re
from io import BytesIO

import fitz  # PyMuPDF


def extract_text_from_file(file_bytes: bytes, file_name: str) -> str:
    """Extract text content from an uploaded file.

    Supports PDFs (with optional OCR for scanned pages), DOCX/DOC documents
    and plain text files.

    Args:
        file_bytes: Raw file data.
        file_name: Original filename used to infer the type.

    Returns:
        Extracted text content, stripped of leading/trailing whitespace.
    """

    if not file_bytes:
        return ""
    file_name = file_name.lower()
    text = ""
    try:
        if file_name.endswith(".pdf"):
            with fitz.open(stream=file_bytes, filetype="pdf") as doc:
                for page in doc:
                    page_text = page.get_text().strip()
                    if page_text:
                        text += page_text + "\n"
                    else:  # fallback to OCR for scanned pages
                        try:
                            import pytesseract
                            from PIL import Image

                            pix = page.get_pixmap()
                            img = Image.open(BytesIO(pix.tobytes("png")))
                            text += pytesseract.image_to_string(img) + "\n"
                        except Exception:
                            pass
        elif file_name.endswith(".docx") or file_name.endswith(".doc"):
            import docx

            doc = docx.Document(BytesIO(file_bytes))
            full_text = [para.text for para in doc.paragraphs]
            text = "\n".join(full_text)
        else:
            text = file_bytes.decode("utf-8", errors="ignore")
    except Exception:
        text = ""
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return "\n".join(lines)


def extract_text_from_url(url: str) -> str:
    """Fetch and clean the main textual content from a web page.

    Args:
        url: Web page URL.

    Returns:
        Cleaned textual content or an empty string if retrieval fails.
    """

    try:
        from readability import Document
        import requests
        from bs4 import BeautifulSoup

        response = requests.get(url, timeout=8)
        if response.status_code != 200:
            return ""
        doc = Document(response.text)
        soup = BeautifulSoup(doc.summary(), "lxml")
        for tag in soup(["script", "style"]):
            tag.decompose()
        text = soup.get_text("\n")
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        return "\n".join(lines)
    except Exception:
        return ""


def merge_texts(*parts: str) -> str:
    """Combine multiple text fragments into one string.

    Empty or whitespace-only parts are ignored. Remaining fragments are
    joined using newline separators.

    Args:
        *parts: Arbitrary text snippets.

    Returns:
        Merged text.
    """

    cleaned = [p.strip() for p in parts if p and p.strip()]
    return "\n".join(cleaned)


def highlight_keywords(text: str, keywords: list[str]) -> str:
    if not text or not keywords:
        return text
    pattern = re.compile(
        "|".join([re.escape(k) for k in keywords]), re.IGNORECASE
    )
    return pattern.sub(lambda m: f"**{m.group(0)}**", text)


def build_boolean_query(job_title: str, skills: list[str]) -> str:
    job_title_part = f"\"{job_title}\"" if job_title else ""
    skill_terms = [f"\"{s.strip()}\"" for s in skills if s.strip()]
    if job_title_part and skill_terms:
        skills_query = " OR ".join(skill_terms)
        return f'{job_title_part} AND ({skills_query})'
    elif job_title_part:
        return job_title_part
    else:
        return " OR ".join(skill_terms)


def seo_optimize(text: str, max_keywords: int = 5) -> dict:
    result = {"keywords": [], "meta_description": ""}
    if not text:
        return result
    words = re.findall(r"\b[a-zA-Z]{4,}\b", text.lower())
    freq = {}
    for w in words:
        freq[w] = freq.get(w, 0) + 1
    top_words = sorted(freq, key=freq.get, reverse=True)
    result["keywords"] = [w for w in top_words[:max_keywords]]
    first_sentence_end = re.search(r'[.!?]', text)
    if first_sentence_end:
        first_sentence = text[:first_sentence_end.end()]
    else:
        first_sentence = text[:160]
    if len(first_sentence) > 160:
        first_sentence = first_sentence[:157] + "..."
    result["meta_description"] = first_sentence.strip()
    return result


def ensure_logs_dir():
    try:
        os.makedirs("logs", exist_ok=True)
    except Exception:
        pass
