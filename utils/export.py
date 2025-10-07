"""Utilities to export generated text into various file formats."""

from __future__ import annotations

from io import BytesIO
import json
from typing import Tuple

import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
import fitz  # PyMuPDF


PDF_FONT_MAP = {
    "Helvetica": "helv",
    "Arial": "helv",
    "Times New Roman": "times",
    "Georgia": "times",
    "Calibri": "helv",
}


def text_to_docx(
    text: str,
    *,
    font: str | None = None,
    logo: bytes | None = None,
    company_name: str | None = None,
) -> bytes:
    """Convert plain text into a DOCX binary with optional styling."""

    doc = docx.Document()

    if logo:
        try:
            image_stream = BytesIO(logo)
            width = Inches(1.8)
            doc.add_picture(image_stream, width=width)
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        except Exception:
            pass

    if company_name:
        heading = doc.add_heading(company_name, level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    normal_style = doc.styles["Normal"]
    if font:
        normal_style.font.name = font

    for line in text.splitlines():
        paragraph = doc.add_paragraph(line)
        if font:
            for run in paragraph.runs:
                run.font.name = font

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def text_to_pdf(
    text: str,
    *,
    font: str | None = None,
    logo: bytes | None = None,
    title: str | None = None,
) -> bytes:
    """Convert text into a single-page PDF respecting basic styling."""

    doc = fitz.open()
    page = doc.new_page()
    top_margin = 72
    left_margin = 72

    if logo:
        try:
            image_rect = fitz.Rect(left_margin, 36, page.rect.width - left_margin, 140)
            page.insert_image(image_rect, stream=logo, keep_proportion=True)
            top_margin = image_rect.y1 + 16
        except Exception:
            top_margin = 72

    if title:
        title_rect = fitz.Rect(
            left_margin, top_margin, page.rect.width - left_margin, top_margin + 40
        )
        page.insert_textbox(
            title_rect,
            title,
            fontsize=18,
            fontname=PDF_FONT_MAP.get(font or "Helvetica", "helv"),
            align=fitz.TEXT_ALIGN_CENTER,
        )
        top_margin = title_rect.y1 + 16

    rect = fitz.Rect(
        left_margin,
        top_margin,
        page.rect.width - left_margin,
        page.rect.height - left_margin,
    )
    page.insert_textbox(
        rect,
        text,
        fontsize=12,
        fontname=PDF_FONT_MAP.get(font or "Helvetica", "helv"),
    )
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def text_to_json(text: str, *, key: str, title: str | None = None) -> bytes:
    """Wrap text in a structured JSON object and return as bytes."""
    data: dict[str, str] = {"type": key, "content": text}
    if title:
        data["title"] = title
    return json.dumps(data, ensure_ascii=False, indent=2).encode("utf-8")


def prepare_download_data(
    text: str,
    fmt: str,
    *,
    key: str,
    title: str | None = None,
    font: str | None = None,
    logo: bytes | None = None,
    company_name: str | None = None,
) -> Tuple[bytes, str, str]:
    """Prepare data, mime type and extension for download."""
    fmt = fmt.lower()
    if fmt == "docx":
        return (
            text_to_docx(text, font=font, logo=logo, company_name=company_name),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "docx",
        )
    if fmt == "pdf":
        return (
            text_to_pdf(text, font=font, logo=logo, title=company_name),
            "application/pdf",
            "pdf",
        )
    if fmt == "json":
        return text_to_json(text, key=key, title=title), "application/json", "json"
    return text.encode("utf-8"), "text/markdown", "md"
