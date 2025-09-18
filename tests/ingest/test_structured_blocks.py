from __future__ import annotations

import io
import sys
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[2]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from docx import Document  # noqa: E402

from ingest.extractors import extract_text_from_file  # noqa: E402
from ingest.reader import clean_structured_document, read_job_text  # noqa: E402


def test_extract_docx_table_blocks() -> None:
    doc = Document()
    doc.add_paragraph("Overview")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Col A"
    table.cell(0, 1).text = "Col B"
    table.cell(1, 0).text = "Row 1"
    table.cell(1, 1).text = "Row 2"
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    setattr(buf, "name", "sample.docx")

    result = clean_structured_document(extract_text_from_file(buf))

    table_blocks = [block for block in result.blocks if block.type == "table"]
    assert table_blocks, "table block should be detected"
    assert table_blocks[0].metadata and table_blocks[0].metadata["rows"][0] == [
        "Col A",
        "Col B",
    ]


def test_read_job_text_detects_bullets() -> None:
    pasted = "- First item\n- Second item"
    doc = read_job_text([], pasted=pasted)
    list_items = [block for block in doc.blocks if block.type == "list_item"]
    assert len(list_items) == 2
    assert [item.text for item in list_items] == ["First item", "Second item"]
