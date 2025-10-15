import json
from io import BytesIO

import docx
import fitz
import pytest

from utils.export import (
    prepare_clean_json,
    prepare_download_data,
    text_to_docx,
    text_to_pdf,
    text_to_json,
)


TEST_LOGO_BYTES = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x08\x00\x00\x00\x08\x08\x02\x00\x00\x00\xed\xc3v?"
    b"\x00\x00\x00\x06bKGD\x00\xff\x00\xff\x00\xff\xa0\xbd\xa7\x93\x00\x00\x00\t pHYs\x00\x00\x0b\x13\x00\x00\x0b\x13\x01\x00\x9a\x9c\x18"
    b"\x00\x00\x00\x07tIME\x07\xe7\x05\x1d\x107\x11\xf0\x82\x98\xa7\x00\x00\x00\x19tEXtComment\x00Created with GIMPd.e\x07\x00\x00\x00>IDAT(\x8d\xad\x93\xb1\x0d\x80 \x0cD3#\xf1\x1d\x12b\x02u"
    b"\x00\x93\xc44>J\x83\xf8\xa1\x17\xef\xe3\xbb\x04\x00\x08\xa6\x9c\x12\xe4\x92\xfc\xaa\xec\xcc\x10\xe0\r\xc7\x1dY\xfbXg\xdf)\x08\x00\x00\x00\x00IEND\xaeB`\x82"
)


def test_text_to_docx() -> None:
    data = text_to_docx("Hello", font="Georgia", company_name="Acme")
    document = docx.Document(BytesIO(data))
    texts = [p.text for p in document.paragraphs if p.text]
    assert "Acme" in texts
    assert "Hello" in texts


def test_text_to_pdf() -> None:
    data = text_to_pdf("Hello")
    doc = fitz.open(stream=data, filetype="pdf")
    assert doc.page_count == 1


def test_text_to_pdf_with_title_and_logo() -> None:
    data = text_to_pdf("Hello", title="Acme Corp", logo=TEST_LOGO_BYTES)
    doc = fitz.open(stream=data, filetype="pdf")
    assert doc.page_count == 1
    page_text = doc.load_page(0).get_text()
    assert "Hello" in page_text
    assert "Acme Corp" in page_text


def test_text_to_json() -> None:
    data = text_to_json("Hello", key="test", title="Title")
    obj = json.loads(data.decode("utf-8"))
    assert obj == {"type": "test", "content": "Hello", "title": "Title"}


def test_prepare_clean_json() -> None:
    payload, mime, ext = prepare_clean_json({"foo": "bär"})
    assert ext == "json"
    assert mime == "application/json"
    assert json.loads(payload.decode("utf-8")) == {"foo": "bär"}


def test_prepare_download_data_docx() -> None:
    data, mime, ext = prepare_download_data(
        "Hello",
        "docx",
        key="test",
        font="Arial",
        company_name="Acme",
    )
    assert ext == "docx"
    assert mime.startswith("application/vnd")
    document = docx.Document(BytesIO(data))
    assert document.paragraphs[0].text == "Acme"


def test_prepare_download_data_docx_includes_logo(monkeypatch: pytest.MonkeyPatch) -> None:
    captured: dict[str, object] = {}

    def fake_text_to_docx(
        text: str,
        *,
        font: str | None = None,
        logo: bytes | None = None,
        company_name: str | None = None,
    ) -> bytes:
        captured["text"] = text
        captured["logo"] = logo
        captured["company_name"] = company_name
        return b"docx"

    monkeypatch.setattr("utils.export.text_to_docx", fake_text_to_docx)

    payload, mime, ext = prepare_download_data(
        "Hello",
        "docx",
        key="test",
        logo=TEST_LOGO_BYTES,
        company_name="Acme",
    )

    assert payload == b"docx"
    assert mime.startswith("application/vnd")
    assert ext == "docx"
    assert captured["logo"] == TEST_LOGO_BYTES
    assert captured["company_name"] == "Acme"


def test_prepare_download_data_pdf_includes_logo(monkeypatch: pytest.MonkeyPatch) -> None:
    captured: dict[str, object] = {}

    def fake_text_to_pdf(
        text: str,
        *,
        font: str | None = None,
        logo: bytes | None = None,
        title: str | None = None,
    ) -> bytes:
        captured["text"] = text
        captured["logo"] = logo
        captured["title"] = title
        return b"pdf"

    monkeypatch.setattr("utils.export.text_to_pdf", fake_text_to_pdf)

    payload, mime, ext = prepare_download_data(
        "Hello",
        "pdf",
        key="test",
        logo=TEST_LOGO_BYTES,
        title="Acme",
    )

    assert payload == b"pdf"
    assert mime == "application/pdf"
    assert ext == "pdf"
    assert captured["logo"] == TEST_LOGO_BYTES
